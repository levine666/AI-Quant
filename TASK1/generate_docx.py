#!/usr/bin/env python3
"""生成 TASK1 Word 文档（宋体、五号、1.5 倍行距、0 段间距、两端对齐）。"""

from __future__ import annotations

import os
import subprocess
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(__file__).resolve().parent
FIG1 = BASE_DIR / "figures" / "fig1_close_price.png"
FIG2 = BASE_DIR / "figures" / "fig2_kline_volume.png"
CSV_PATH = BASE_DIR / "data" / "smic_688981_daily.csv"

STUDENT_NAME = os.environ.get("STUDENT_NAME", "姓名")
OUTPUT_DOCX = BASE_DIR / f"{STUDENT_NAME}TASK1.docx"

FONT_NAME = "宋体"
BODY_SIZE = Pt(10.5)  # 五号
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)


def ensure_data() -> None:
    if not CSV_PATH.exists() or not FIG1.exists():
        subprocess.run([sys.executable, str(BASE_DIR / "akshare_fetch.py")], check=True)


def set_run_font(run, size: Pt, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    run.font.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: Pt = BODY_SIZE,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent: bool = True,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 约 2 字符
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    size = H1_SIZE if level == 1 else H2_SIZE
    add_paragraph(
        doc,
        text,
        size=size,
        bold=True,
        indent=False,
        space_before=Pt(6),
        space_after=Pt(6),
    )


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 15.0) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_paragraph(doc, caption, size=H2_SIZE, bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)


def load_stats() -> tuple[int, str, str]:
    if not CSV_PATH.exists():
        return 0, "—", "—"
    df = pd.read_csv(CSV_PATH)
    rows = len(df)
    date_range = f"{df['date'].iloc[0]} ~ {df['date'].iloc[-1]}"
    ret = (df["close"].iloc[-1] / df["close"].iloc[0] - 1) * 100
    return rows, date_range, f"{ret:+.2f}%"


def build_document() -> Document:
    rows, date_range, ret_text = load_stats()
    doc = Document()

    # 默认正文样式
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = BODY_SIZE

    add_paragraph(doc, "量化交易基础与 AkShare 数据实践", size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=Pt(6))
    add_paragraph(doc, "TASK1 作业报告", size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=Pt(12))

    # ===== 问题 1 =====
    add_heading(doc, "一、量化交易相较于传统手工交易的优势", 1)
    for text in [
        "随着金融市场数据化、自动化程度不断提高，量化交易（Quantitative Trading）已成为机构与个人投资者的重要参与方式。相较于依赖人工盯盘、凭经验下单的传统手工交易，量化交易在效率、纪律性与可验证性方面具有显著优势。",
        "第一，量化交易能够克服人性弱点。手工交易容易受到贪婪、恐惧、过度自信等情绪影响，导致追涨杀跌、频繁交易等非理性行为。量化策略依据预先设定的规则自动执行，减少主观情绪对决策的干扰。",
        "第二，量化交易具有更高的执行效率。计算机可在毫秒级完成行情扫描、信号识别与下单，能同时监控大量标的，捕捉短暂出现的交易机会。",
        "第三，量化交易支持历史回测与策略验证。策略可在多年历史数据上模拟运行，评估收益率、最大回撤等指标，在实盘前检验有效性。",
        "第四，量化交易强调一致性与可复制性。同一套规则可重复应用，便于团队协作与策略迭代。",
        "第五，量化交易便于风险管控。可通过程序设定止损、止盈、仓位上限等规则，实时监控风险敞口。",
        "第六，量化交易能够处理海量信息。量化模型可综合价格、成交量、财务等多维数据进行分析，而人工分析的信息容量有限。",
        "综上所述，量化交易将投资逻辑代码化、流程化、数据化，为策略研究提供了可量化、可复现的框架。本次作业通过 Python + AkShare 获取行情、绘图、存档，正是量化研究流程的具体实践。",
    ]:
        add_paragraph(doc, text)

    doc.add_page_break()

    # ===== 问题 2 =====
    add_heading(doc, "二、基本概念解释：K 线、基本面、技术面", 1)
    add_heading(doc, "（一）K 线", 2)
    for text in [
        "K 线（Candlestick Chart）是展示证券价格在某一时间段内开盘价、收盘价、最高价、最低价及其涨跌关系的图形工具。每根 K 线由实体与上下影线组成：实体上端为收盘价、下端为开盘价（A 股习惯红色上涨、绿色下跌）；影线表示最高价与最低价。",
        "K 线可按日 K、周 K、月 K 等周期划分。通过连续 K 线组合，可观察趋势、支撑/压力位及常见形态，是技术分析中最基础的可视化方式。本作业图2 即采用日 K 线展示中芯国际价格与成交量。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "（二）基本面", 2)
    for text in [
        "基本面（Fundamental Analysis）从宏观经济、行业状况与公司内在价值出发，评估证券合理价格。其核心逻辑是价格长期围绕内在价值波动，通过研究盈利能力与成长性，判断资产是否被高估或低估。",
        "股票基本面包括：宏观层面（GDP、利率、政策等）；行业层面（竞争格局、景气周期）；公司层面（营收、净利润、现金流、市盈率等）。基本面分析适合中长期投资，回答“买什么、值不值得买”。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "（三）技术面", 2)
    for text in [
        "技术面（Technical Analysis）基于历史价格、成交量等市场行为数据，运用图表与指标预测价格走势。其假设包括：市场行为包容一切信息、价格沿趋势运动、历史会重演。",
        "常见工具有 K 线形态、移动平均线（MA）、MACD、RSI 等。技术面侧重“何时买卖”，与基本面形成互补。本作业通过收盘价曲线（图1）与日 K 线（图2）进行技术面观察。",
    ]:
        add_paragraph(doc, text)

    doc.add_page_break()

    # ===== 问题 3 =====
    add_heading(doc, "三、AkShare 数据获取与 Python 实现", 1)
    add_heading(doc, "（一）AkShare 平台与数据接口", 2)
    for text in [
        "AkShare（https://akshare.akfamily.xyz/）是开源 Python 财经数据接口库，pip install akshare 即可使用，无需注册 Token，适合量化入门与教学实验。",
        "本作业以半导体龙头中芯国际 A 股（688981）为例，调用 ak.stock_zh_a_daily 接口获取上交所科创板股票近一年每个交易日的日线数据（前复权），字段包括 date、open、high、low、close、volume、amount 等。",
        "在完成作业基本要求之外，我们还使用 ak.stock_hk_daily 获取港股 00981 数据，用 ak.currency_boc_sina 获取港元/人民币汇率，完成了 A/H 股对比分析与 GitHub Pages 交互看板（dashboard.html），详见 smic_analysis.ipynb。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "（二）Python 实现步骤", 2)
    for text in [
        "实现流程如下：① 安装 akshare、pandas、matplotlib；② 设定股票代码 sh688981 与起止日期（近 365 个自然日）；③ 调用 stock_zh_a_daily 拉取日线并按 date 排序；④ 用 matplotlib 绘制每日收盘价曲线（图1）；⑤ 用 mplfinance 绘制日 K 线与成交量（图2）；⑥ 将数据保存为 UTF-8 编码 CSV，供后续任务使用。完整代码见 TASK1/akshare_fetch.py。",
        "核心代码片段：df = ak.stock_zh_a_daily(symbol='sh688981', start_date='...', end_date='...', adjust='qfq')。adjust='qfq' 表示前复权，消除除权除息对价格连续性的影响，便于趋势分析。",
    ]:
        add_paragraph(doc, text)

    add_figure(doc, FIG1, "图1  中芯国际（688981）近一年每日收盘价")
    for text in [
        f"图1 展示了中芯国际（688981）在过去约 {rows} 个交易日内的前复权收盘价走势（{date_range}）。解读如下：① 区间整体涨跌幅约为 {ret_text}，反映市场对公司价值的重估过程；② 2025 年下半年至 2026 年上半年价格中枢明显抬升，显示资金关注度提高；③ 近期出现回调，需结合成交量（图2）判断是技术性调整还是趋势反转。",
        "收盘价曲线是最常用的技术面观察工具，可快速识别趋势方向与关键价位。相比手工在 Excel 中逐日录入，Python + AkShare 可一键获取完整历史数据并自动出图，体现了量化研究在效率上的优势。",
    ]:
        add_paragraph(doc, text)

    add_figure(doc, FIG2, "图2  中芯国际（688981）日K线及成交量")
    for text in [
        "图2 以日 K 线形式展示开高低收，下方柱状图为成交量。解读如下：① 红色 K 线表示当日收盘价高于开盘价（上涨），绿色表示下跌，符合 A 股配色习惯；② 价涨量增通常表示上涨动能较强，价跌量缩可能表示抛压减轻；③ 图中可见若干放量长阳，往往对应重要信息释放或资金集中流入的时段。",
        "该图与项目 dashboard.html 看板、smic_analysis.ipynb Notebook 中的 K 线分析一致，是从数据获取到可视化展示的完整技术面实践。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "（三）CSV 数据保存", 2)
    for text in [
        f"数据已保存至 TASK1/data/smic_688981_daily.csv，共 {rows} 行，字段包括 date、open、high、low、close、volume、amount、outstanding_share、turnover 等，可直接用 pandas.read_csv 读取。",
        "CSV 采用 utf-8-sig 编码，兼容 Excel 中文显示。项目根目录 data/ 下还保存了 A/H 股对比数据（smic_688981_a.csv、smic_00981_hk.csv），可用于后续因子计算、回测或 AH 溢价分析。",
    ]:
        add_paragraph(doc, text)

    return doc


def main() -> None:
    ensure_data()
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"Word 已生成: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
